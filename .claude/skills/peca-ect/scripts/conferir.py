# -*- coding: utf-8 -*-
"""Confere uma peça gerada contra o _FORMATO_BASE.docx. Uso: python3 conferir.py peca.docx"""
import hashlib, re, sys, zipfile
from peca_fmt import BASE

ALVO = sys.argv[1]
PARTES = ["word/header1.xml", "word/footer1.xml", "word/styles.xml",
          "word/media/image1.png", "word/settings.xml"]

z, b = zipfile.ZipFile(ALVO), zipfile.ZipFile(BASE)
erros = []
for parte in PARTES:
    if hashlib.md5(z.read(parte)).hexdigest() != hashlib.md5(b.read(parte)).hexdigest():
        erros.append(f"formatação-base alterada em {parte}")
print("formatação-base preservada" if not erros else "FALHA na formatação-base")

try:
    import docx
except ImportError:
    sys.exit("instale python-docx para a conferência de conteúdo")

d = docx.Document(ALVO)
s = d.sections[0]
if round(s.left_margin.cm, 1) != 3.0 or round(s.right_margin.cm, 1) != 2.0:
    erros.append("margens fora do padrão")

papeis = {"retângulo": 0, "subtópico/alínea": 0, "citação": 0, "corpo": 0}
for p in d.paragraphs:
    if not p.text.strip():
        continue
    pf, xml = p.paragraph_format, p._p.xml
    tam = {r.font.size.pt for r in p.runs if r.font.size}
    esq = None if pf.left_indent is None else round(pf.left_indent.cm, 1)
    pri = None if pf.first_line_indent is None else round(pf.first_line_indent.cm, 1)
    if "<w:pBdr>" in xml:
        papeis["retângulo"] += 1
        if any(r.underline for r in p.runs):
            erros.append(f"tópico em retângulo sublinhado (não deve ser): {p.text[:50]}")
    elif 10.0 in tam:
        papeis["citação"] += 1
    elif esq == 3.0:
        papeis["subtópico/alínea"] += 1
    elif pri == 3.0:
        papeis["corpo"] += 1
    limpo = re.sub(r"\(\.\.\.\)\s*", "", p.text)
    for ruim in ["  ", " ,", ", ,", ",,"]:
        if ruim in limpo:
            erros.append(f"tipografia {ruim!r}: {p.text[:60]}")

print("papéis:", ", ".join(f"{k}={v}" for k, v in papeis.items()))
texto = "\n".join(p.text for p in d.paragraphs)
print("marcações [REVISAR]:", texto.count("[REVISAR"))
if erros:
    print("\nPROBLEMAS:")
    for e in erros:
        print("  -", e)
    sys.exit(1)
print("OK")
